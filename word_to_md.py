import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsmap
from docx.opc.constants import RELATIONSHIP_TYPE as RT


def get_run_text_with_formatting(run):
    """提取 run 的文本并应用 Markdown 格式"""
    text = run.text
    if not text:
        return ''

    # 检测格式
    is_bold = run.bold
    is_italic = run.italic
    is_code = False

    # 检测是否为代码字体
    if run.font.name and 'Courier' in run.font.name:
        is_code = True

    # 应用格式
    if is_code:
        text = f'`{text}`'
    else:
        if is_bold and is_italic:
            text = f'***{text}***'
        elif is_bold:
            text = f'**{text}**'
        elif is_italic:
            text = f'*{text}*'

    return text


def extract_inline_formatting(paragraph):
    """提取段落中所有 run 的格式化文本"""
    result = []
    for run in paragraph.runs:
        formatted_text = get_run_text_with_formatting(run)
        if formatted_text:
            result.append(formatted_text)
    return ''.join(result)


def is_code_paragraph(paragraph):
    """检测段落是否为代码块"""
    # 检查样式名称
    style_name = paragraph.style.name
    if style_name == 'Code':
        return True

    # 检查段落格式是否有缩进（代码块特征）
    pf = paragraph.paragraph_format
    has_indent = pf.left_indent and pf.left_indent > Inches(0)

    # 检查所有 run 是否使用 Courier 字体且字体大小为 9pt
    all_courier = True
    has_9pt = False
    for run in paragraph.runs:
        if run.text.strip():
            font_name = run.font.name
            font_size = run.font.size
            if not (font_name and 'Courier' in font_name):
                all_courier = False
            if font_size and font_size == Pt(9):
                has_9pt = True

    # 只有同时满足 Courier 字体和缩进才认为是代码块
    # 或者样式为 Code
    return (all_courier and has_indent) or (all_courier and has_9pt)


def convert_table_to_md(table):
    """将 Word 表格转换为 Markdown 表格"""
    rows = table.rows
    if not rows:
        return ''

    lines = []

    # 处理表头（第一行）
    header_cells = rows[0].cells
    header_texts = [cell.text.strip() for cell in header_cells]
    lines.append('| ' + ' | '.join(header_texts) + ' |')

    # 添加分隔行
    separator = '|' + '|'.join(['---' for _ in header_cells]) + '|'
    lines.append(separator)

    # 处理数据行
    for row in rows[1:]:
        cells = row.cells
        cell_texts = [cell.text.strip() for cell in cells]
        lines.append('| ' + ' | '.join(cell_texts) + ' |')

    return '\n'.join(lines)


def extract_images(doc, output_dir, md_file_name):
    """提取文档中的图片并保存到指定目录"""
    image_map = {}  # rId -> 图片文件名
    image_counter = 0

    # 确保 images 目录存在
    images_dir = os.path.join(output_dir, 'images')
    if not os.path.exists(images_dir):
        os.makedirs(images_dir)

    # 遍历文档中的所有关系，查找图片
    for rel in doc.part.rels.values():
        if rel.reltype == RT.IMAGE:
            image_counter += 1
            image_data = rel.target_part.blob

            # 根据内容类型确定扩展名
            content_type = rel.target_part.content_type
            ext_map = {
                'image/png': '.png',
                'image/jpeg': '.jpg',
                'image/jpg': '.jpg',
                'image/gif': '.gif',
                'image/bmp': '.bmp',
                'image/tiff': '.tiff',
            }
            ext = ext_map.get(content_type, '.png')

            # 生成图片文件名
            base_name = os.path.splitext(md_file_name)[0]
            image_name = f'{base_name}_image_{image_counter}{ext}'
            image_path = os.path.join(images_dir, image_name)

            # 保存图片
            with open(image_path, 'wb') as f:
                f.write(image_data)

            image_map[rel.rId] = f'images/{image_name}'

    return image_map


def get_paragraph_image(paragraph, image_map):
    """检查段落是否包含图片，返回图片路径"""
    # 检查段落中的绘图元素
    for run in paragraph.runs:
        # 检查 run 中的绘图元素
        drawing_elements = run._element.findall('.//a:graphic', namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
        for graphic in drawing_elements:
            # 查找图片引用
            blip = graphic.find('.//a:blip', namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
            if blip is not None:
                # 获取嵌入图片的 rId
                embed_attr = blip.get(qn('r:embed'))
                if embed_attr and embed_attr in image_map:
                    return image_map[embed_attr]

    # 另一种方式：通过 inline 元素查找
    for run in paragraph.runs:
        inline_elements = run._element.findall('.//wp:inline', namespaces={'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'})
        for inline in inline_elements:
            graphic = inline.find('.//a:graphic', namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
            if graphic is not None:
                blip = graphic.find('.//a:blip', namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
                if blip is not None:
                    embed_attr = blip.get(qn('r:embed'))
                    if embed_attr and embed_attr in image_map:
                        return image_map[embed_attr]

    return None


def is_quote_paragraph(paragraph):
    """检测段落是否为引用块"""
    style_name = paragraph.style.name
    return 'Quote' in style_name or style_name == 'Intense Quote'


def get_list_level(paragraph):
    """获取列表的嵌套级别"""
    # 通过段落缩进判断嵌套级别
    pf = paragraph.paragraph_format
    if pf.left_indent:
        # 每级缩进约 0.25 英寸
        level = int(pf.left_indent / Inches(0.25))
        return max(0, level - 1)
    return 0


def word_to_md(docx_file, output_md_file):
    """将 Word 文档转换为 Markdown"""
    # 加载 Word 文档
    doc = Document(docx_file)

    # 获取输出目录
    output_dir = os.path.dirname(output_md_file)
    md_file_name = os.path.basename(output_md_file)

    # 提取图片
    image_map = extract_images(doc, output_dir, md_file_name)

    # 收集转换后的 Markdown 行
    md_lines = []

    # 用于跟踪代码块
    in_code_block = False
    code_lines = []

    # 用于跟踪列表
    in_list = False

    # 遍历文档元素
    for element in doc.element.body:
        # 处理段落
        if element.tag == qn('w:p'):
            # 找到对应的段落对象
            paragraph = None
            for p in doc.paragraphs:
                if p._element == element:
                    paragraph = p
                    break

            if paragraph is None:
                continue

            # 检查是否包含图片
            image_path = get_paragraph_image(paragraph, image_map)
            if image_path:
                md_lines.append(f'![图片]({image_path})')
                continue

            # 检查标题
            style_name = paragraph.style.name
            if style_name.startswith('Heading'):
                # 结束代码块（如果有）
                if in_code_block:
                    md_lines.append('\n```\n')
                    in_code_block = False

                try:
                    level = int(style_name.split()[-1])
                    text = paragraph.text.strip()
                    if text:
                        md_lines.append(f'\n{"#" * level} {text}\n')
                except ValueError:
                    # 无法解析级别，作为普通段落处理
                    text = extract_inline_formatting(paragraph)
                    if text.strip():
                        md_lines.append(text)
                continue

            # 检查列表样式
            if 'List' in style_name:
                # 结束代码块（如果有）
                if in_code_block:
                    md_lines.append('\n```\n')
                    in_code_block = False

                text = paragraph.text.strip()
                if text:
                    level = get_list_level(paragraph)
                    indent = '    ' * level  # 每级 4 个空格

                    if 'Bullet' in style_name:
                        md_lines.append(f'{indent}- {text}')
                    elif 'Number' in style_name:
                        # 尝试提取编号
                        md_lines.append(f'{indent}1. {text}')
                    else:
                        md_lines.append(f'{indent}- {text}')
                continue

            # 检查引用块
            if is_quote_paragraph(paragraph):
                # 结束代码块（如果有）
                if in_code_block:
                    md_lines.append('\n```\n')
                    in_code_block = False

                text = paragraph.text.strip()
                if text:
                    md_lines.append(f'> {text}')
                continue

            # 检查代码块
            if is_code_paragraph(paragraph):
                if not in_code_block:
                    md_lines.append('\n```')
                    in_code_block = True
                text = paragraph.text
                # 不 strip，保留缩进
                md_lines.append(text)
                continue

            # 结束代码块（如果有）
            if in_code_block:
                md_lines.append('\n```\n')
                in_code_block = False

            # 普通段落
            text = extract_inline_formatting(paragraph)
            if text.strip():
                md_lines.append(text)

        # 处理表格
        elif element.tag == qn('w:tbl'):
            # 结束代码块（如果有）
            if in_code_block:
                md_lines.append('\n```\n')
                in_code_block = False

            # 找到对应的表格对象
            table = None
            for t in doc.tables:
                if t._element == element:
                    table = t
                    break

            if table is not None:
                md_table = convert_table_to_md(table)
                md_lines.append('\n' + md_table + '\n')

    # 结束最后的代码块（如果有）
    if in_code_block:
        md_lines.append('\n```\n')

    # 合并所有行
    md_content = '\n'.join(md_lines)

    # 清理多余的空行
    md_content = re.sub(r'\n{3,}', '\n\n', md_content)

    # 保存 Markdown 文件
    with open(output_md_file, 'w', encoding='utf-8') as f:
        f.write(md_content)

    print(f'转换完成: {output_md_file}')


if __name__ == '__main__':
    # 转换 WordToMd 文件夹中的所有 docx 文件
    md_folder = 'WordToMd'

    # 检查 WordToMd 文件夹是否存在
    if not os.path.exists(md_folder):
        print(f"文件夹 '{md_folder}' 不存在，已创建")
        os.makedirs(md_folder)
        os.makedirs(os.path.join(md_folder, 'images'))

    # 获取 WordToMd 文件夹中的所有 docx 文件
    docx_files = [f for f in os.listdir(md_folder) if f.endswith('.docx')]

    if not docx_files:
        print(f"在 '{md_folder}' 文件夹中没有找到 .docx 文件")
        exit(0)

    for docx_file in docx_files:
        docx_path = os.path.join(md_folder, docx_file)
        output_file = docx_file.replace('.docx', '.md')
        output_path = os.path.join(md_folder, output_file)
        word_to_md(docx_path, output_path)