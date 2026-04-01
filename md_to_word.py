import os
import re
import markdown
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def parse_inline_formatting(paragraph, text):
    parts = re.split(r'(<[^>]+>)', text)
    current_run = None
    bold = False
    italic = False
    code = False
    
    for part in parts:
        if part.startswith('<strong>') or part.startswith('<b>'):
            bold = True
        elif part.startswith('</strong>') or part.startswith('</b>'):
            bold = False
        elif part.startswith('<em>') or part.startswith('<i>'):
            italic = True
        elif part.startswith('</em>') or part.startswith('</i>'):
            italic = False
        elif part.startswith('<code>'):
            code = True
        elif part.startswith('</code>'):
            code = False
        elif part and not part.startswith('<'):
            if code:
                run = paragraph.add_run(part)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
            else:
                run = paragraph.add_run(part)
                if bold:
                    run.bold = True
                if italic:
                    run.italic = True

def md_to_word(md_file, output_file):
    # 读取markdown文件
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 转换为HTML，启用更多扩展
    html = markdown.markdown(md_content, extensions=['tables', 'fenced_code', 'codehilite', 'extra'])
    
    # 创建Word文档
    doc = Document()
    
    # 添加代码样式
    styles = doc.styles
    code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Courier New'
    code_style.font.size = Pt(9)
    code_style.font.color.rgb = None
    code_style.paragraph_format.left_indent = Inches(0.25)
    code_style.paragraph_format.right_indent = Inches(0.25)
    code_style.paragraph_format.space_before = Pt(6)
    code_style.paragraph_format.space_after = Pt(6)
    
    # 设置中文字体
    code_style.font._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
    
    # 处理内容
    lines = html.split('\n')
    current_paragraph = None
    in_code_block = False
    code_content = []
    in_table = False
    table = None
    current_row = None
    current_cell = 0
    table_cols = 0
    in_list = False
    list_type = None
    list_level = 0
    list_counter = 0
    pending_paragraph_text = None
    current_paragraph_html = None
    list_stack = []
    
    for line in lines:
        # 对于代码块，不使用 strip()，保留缩进
        if in_code_block:
            if not line:
                continue
        else:
            line = line.strip()
            if not line:
                continue
        
        # 处理标题
        if line.startswith('<h1>'):
            text = line.replace('<h1>', '').replace('</h1>', '')
            doc.add_heading(text, level=1)
        elif line.startswith('<h2>'):
            text = line.replace('<h2>', '').replace('</h2>', '')
            doc.add_heading(text, level=2)
        elif line.startswith('<h3>'):
            text = line.replace('<h3>', '').replace('</h3>', '')
            doc.add_heading(text, level=3)
        # 处理代码块
        elif 'codehilite' in line or '<pre><code' in line:
            in_code_block = True
            code_content = []
            # 处理当前行中的内容
            clean_line = re.sub(r'<[^>]+>', '', line)
            if clean_line:
                code_content.append(clean_line)
        elif '</code></pre>' in line:
            in_code_block = False
            if code_content:
                # 添加代码段落，解码HTML实体
                decoded_content = []
                for line in code_content:
                    # 解码HTML实体
                    line = line.replace('&quot;', '"').replace('&lt;', '<').replace('&gt;', '>').replace('&amp;', '&')
                    decoded_content.append(line)
                code_paragraph = doc.add_paragraph('\n'.join(decoded_content), style='Code')
                code_paragraph.paragraph_format.space_before = Pt(12)
                code_paragraph.paragraph_format.space_after = Pt(12)
        elif in_code_block:
            # 移除代码高亮的HTML标签
            clean_line = re.sub(r'<[^>]+>', '', line)
            if clean_line:  # 只添加非空行
                code_content.append(clean_line)
        # 处理独立的图片标签
        elif line.startswith('<img'):
            img_match = re.search(r'<img[^>]*src="([^"]+)"[^>]*>', line)
            if img_match:
                img_path = img_match.group(1)
                # 处理相对路径
                if not os.path.isabs(img_path):
                    img_path = os.path.join(os.path.dirname(md_file), img_path)
                
                # 处理反斜杠路径
                img_path = img_path.replace('\\', '/')
                
                if os.path.exists(img_path):
                    # 添加图片
                    doc.add_picture(img_path, width=Inches(5))
                    # 居中图片
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 处理段落
        elif line.startswith('<p>'):
            # 如果已经有未完成的段落，先处理它
            if current_paragraph_html is not None:
                text = current_paragraph_html.replace('<p>', '').replace('</p>', '')
                if text:
                    para = doc.add_paragraph()
                    parse_inline_formatting(para, text)
                current_paragraph_html = None
            
            # 检查是否是段落开始（包含 </p>）
            if '</p>' in line:
                text = line.replace('<p>', '').replace('</p>', '')
                
                # 检查是否包含图片
                img_match = re.search(r'<img src="([^"]+)"[^>]*>', text)
                if img_match:
                    img_path = img_match.group(1)
                    # 处理相对路径
                    if not os.path.isabs(img_path):
                        img_path = os.path.join(os.path.dirname(md_file), img_path)
                    
                    # 处理反斜杠路径
                    img_path = img_path.replace('\\', '/')
                    
                    if os.path.exists(img_path):
                        # 添加图片
                        doc.add_picture(img_path, width=Inches(5))
                        # 居中图片
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # 移除图片标签
                    text = re.sub(r'<img src="[^"]+"[^>]*>', '', text)
                
                # 检查是否包含列表项（以 - 开头）
                list_items = re.findall(r'[\s\n](-\s+[^\n]+)', text)
                
                if list_items:
                    # 提取段落文本（移除列表项）
                    paragraph_text = re.sub(r'[\s\n]-\s+[^\n]+', '', text).strip()
                    
                    # 添加段落文本
                    if paragraph_text:
                        para = doc.add_paragraph()
                        parse_inline_formatting(para, paragraph_text)
                    
                    # 添加列表
                    for item in list_items:
                        item_text = item.strip()
                        if item_text.startswith('- '):
                            item_text = item_text[2:]
                        para = doc.add_paragraph(style='List Bullet')
                        parse_inline_formatting(para, item_text)
                else:
                    if text:
                        para = doc.add_paragraph()
                        parse_inline_formatting(para, text)
            else:
                # 段落开始但没有结束，保存到变量中
                current_paragraph_html = line
        # 处理段落续行（不以 <p> 开头但包含 HTML 标签）
        elif current_paragraph_html is not None and ('<' in line and '>' in line):
            # 追加到当前段落
            current_paragraph_html += ' ' + line
            # 如果包含 </p>，处理整个段落
            if '</p>' in line:
                text = current_paragraph_html.replace('<p>', '').replace('</p>', '')
                if text:
                    para = doc.add_paragraph()
                    parse_inline_formatting(para, text)
                current_paragraph_html = None
        # 处理独立的列表项（不在 <p> 标签内）
        elif line.startswith('- ') or (line.startswith('-') and len(line) > 1 and line[1] == ' '):
            item_text = line[2:] if line.startswith('- ') else line[1:].strip()
            
            # 如果有待处理的段落文本，先添加它
            if pending_paragraph_text:
                para = doc.add_paragraph()
                parse_inline_formatting(para, pending_paragraph_text)
                pending_paragraph_text = None
            
            # 添加列表项
            para = doc.add_paragraph(style='List Bullet')
            parse_inline_formatting(para, item_text)
        # 处理以 - 开头但后面没有空格的情况（可能是列表项的结束标签）
        elif line.startswith('-') and line.endswith('</p>'):
            item_text = line[1:-4].strip()
            
            # 如果有待处理的段落文本，先添加它
            if pending_paragraph_text:
                para = doc.add_paragraph()
                parse_inline_formatting(para, pending_paragraph_text)
                pending_paragraph_text = None
            
            # 添加列表项
            para = doc.add_paragraph(style='List Bullet')
            parse_inline_formatting(para, item_text)
        # 处理列表
        elif line.startswith('<ul>'):
            in_list = True
            list_type = 'ul'
            list_counter = 0
            list_level = len(list_stack)
            list_stack.append(list_type)
        elif line.startswith('<ol>'):
            in_list = True
            list_type = 'ol'
            list_counter = 0
            list_level = len(list_stack)
            list_stack.append(list_type)
        elif line.startswith('</ul>') or line.startswith('</ol>'):
            in_list = False
            if list_stack:
                list_stack.pop()
            list_type = list_stack[-1] if list_stack else None
            list_level = len(list_stack)
            current_paragraph = None
        elif line.startswith('<li>') and in_list:
            text = line.replace('<li>', '').replace('</li>', '')
            text = text.strip()
            
            if list_type == 'ul':
                para = doc.add_paragraph(style='List Bullet')
            elif list_type == 'ol':
                para = doc.add_paragraph(style='List Number')
            
            # 根据嵌套级别设置缩进
            if list_level > 0:
                para.paragraph_format.left_indent = Inches(0.5 * (list_level + 1))
                para.paragraph_format.first_line_indent = Inches(-0.25)
            
            parse_inline_formatting(para, text)
            current_paragraph = para
        # 处理表格
        elif line.startswith('<table>'):
            in_table = True
            table = doc.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            current_row = table.rows[0]
            current_cell = 0
            table_cols = 0
        elif line.startswith('</table>'):
            in_table = False
            table = None
            current_row = None
            current_cell = 0
            table_cols = 0
        elif line.startswith('<thead>') and in_table:
            continue
        elif line.startswith('</thead>') and in_table:
            continue
        elif line.startswith('<tbody>') and in_table:
            continue
        elif line.startswith('</tbody>') and in_table:
            continue
        elif line.startswith('<tr>') and in_table:
            if table and current_row:
                current_cell = 0
            else:
                table.add_row()
                current_row = table.rows[-1]
                current_cell = 0
        elif line.startswith('</tr>'):
            if table and not current_row:
                table.add_row()
            # 确保表格有足够的列
            if table and len(table.columns) < table_cols:
                while len(table.columns) < table_cols:
                    table.add_column(Inches(2))
            current_row = None
        elif line.startswith('<th>') and in_table:
            text = line.replace('<th>', '').replace('</th>', '')
            if current_row:
                if current_cell >= len(current_row.cells):
                    table.add_column(Inches(2))
                current_row.cells[current_cell].text = text
                current_row.cells[current_cell].paragraphs[0].runs[0].bold = True
                current_cell += 1
                table_cols = max(table_cols, current_cell)
        elif line.startswith('<td>') and in_table:
            text = line.replace('<td>', '').replace('</td>', '')
            if current_row:
                if current_cell >= len(current_row.cells):
                    table.add_column(Inches(2))
                current_row.cells[current_cell].text = text
                current_cell += 1
    
    # 处理待处理的段落文本（如果有）
    if pending_paragraph_text:
        para = doc.add_paragraph()
        parse_inline_formatting(para, pending_paragraph_text)
        pending_paragraph_text = None
    
    # 保存文档
    # 如果文件已存在，先删除
    if os.path.exists(output_file):
        try:
            os.remove(output_file)
        except Exception as e:
            print(f"警告: 无法删除现有文件 {output_file}: {e}")
            # 尝试使用临时文件名
            output_file = output_file.replace('.docx', '_new.docx')
    
    doc.save(output_file)
    print(f"转换完成: {output_file}")

if __name__ == "__main__":
    # 转换 MdToWord 文件夹中的所有 md 文件
    doc_folder = 'MdToWord'
    
    # 检查 MdToWord 文件夹是否存在
    if not os.path.exists(doc_folder):
        print(f"错误: 文件夹 '{doc_folder}' 不存在")
        exit(1)
    
    # 获取 MdToWord 文件夹中的所有 md 文件
    md_files = [f for f in os.listdir(doc_folder) if f.endswith('.md')]
    
    if not md_files:
        print(f"在 '{doc_folder}' 文件夹中没有找到 .md 文件")
        exit(0)
    
    for md_file in md_files:
        md_path = os.path.join(doc_folder, md_file)
        output_file = md_file.replace('.md', '.docx')
        output_path = os.path.join(doc_folder, output_file)
        md_to_word(md_path, output_path)
